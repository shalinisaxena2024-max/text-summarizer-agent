"""
tools.py — Input Loading Utilities
===================================
Provides functions to load text from various sources:
- Plain text strings
- Web URLs
- Local files (.txt, .pdf, .docx)
"""

import logging
from pathlib import Path

import requests
from PyPDF2 import PdfReader
from docx import Document

log = logging.getLogger(__name__)


def load_plain_text(text: str) -> str:
    """
    Load plain text directly from input.
    
    Args:
        text: Raw text string.
    
    Returns:
        The input text as-is.
    """
    if not text or not text.strip():
        raise ValueError("Plain text input is empty.")
    return text.strip()


def load_from_url(url: str) -> str:
    """
    Fetch and extract text content from a URL.
    
    Args:
        url: Web URL to scrape.
    
    Returns:
        Extracted text content from the page.
    
    Raises:
        requests.RequestException: If the URL fetch fails.
        ValueError: If no text content is found.
    """
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Try to extract text using basic HTML parsing
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
        except ImportError:
            # Fallback if BeautifulSoup not installed
            text = response.text
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        text = '\n'.join(line for line in lines if line)
        
        if not text.strip():
            raise ValueError(f"No text content extracted from {url}")
        
        log.info("Successfully fetched content from URL (%d chars)", len(text))
        return text.strip()
    
    except requests.RequestException as e:
        raise requests.RequestException(f"Failed to fetch URL {url}: {e}") from e


def load_from_file(file_path: str) -> str:
    """
    Load text from a local file (.txt, .pdf, or .docx).
    
    Args:
        file_path: Path to the file.
    
    Returns:
        Extracted text content from the file.
    
    Raises:
        FileNotFoundError: If the file doesn't exist.
        ValueError: If the file format is not supported or file is empty.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if path.suffix.lower() == ".txt":
        text = path.read_text(encoding='utf-8')
    
    elif path.suffix.lower() == ".pdf":
        try:
            text = _extract_pdf_text(str(path))
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}") from e
    
    elif path.suffix.lower() == ".docx":
        try:
            text = _extract_docx_text(str(path))
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}") from e
    
    else:
        raise ValueError(
            f"Unsupported file format: {path.suffix}. "
            "Supported formats: .txt, .pdf, .docx"
        )
    
    if not text.strip():
        raise ValueError(f"File is empty or contains no extractable text: {file_path}")
    
    log.info("Successfully loaded file %s (%d chars)", path.name, len(text))
    return text.strip()


def _extract_pdf_text(pdf_path: str) -> str:
    """Extract text from a PDF file."""
    text_parts = []
    try:
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text_parts.append(page.extract_text())
                except Exception as e:
                    log.warning("Could not extract page %d: %s", page_num, e)
    except Exception as e:
        raise ValueError(f"Error reading PDF: {e}") from e
    
    return '\n'.join(text_parts)


def _extract_docx_text(docx_path: str) -> str:
    """Extract text from a DOCX file."""
    text_parts = []
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {e}") from e
    
    return '\n'.join(text_parts)
